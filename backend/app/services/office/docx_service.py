import logging
from docx import Document

logger = logging.getLogger(__name__)


def _resolve_blocks(json_content: dict) -> list:
    """Extract a flat list of block dicts from DocumentIR or legacy format."""
    if "blocks" in json_content:
        return json_content["blocks"]
    content = json_content.get("content", json_content) if isinstance(json_content, dict) else json_content
    return content if isinstance(content, list) else []


class DocxService:

    async def export(self, file_path: str, json_content: dict) -> None:
        doc = Document()
        blocks = _resolve_blocks(json_content)

        for item in blocks:
            bt = item.get("type", "")
            text = item.get("text", "")

            if bt in ("paragraph", "段落"):
                doc.add_paragraph(text)
            elif bt == "heading":
                level = item.get("level", 1)
                doc.add_heading(text, level=min(level, 9))
            elif bt == "table":
                rows = item.get("rows", [])
                if text and not rows:
                    rows = [{"cells": [c.strip() for c in row.split("|")]}
                            for row in text.split("\n") if "|" in row]
                if rows:
                    first_row = rows[0].get("cells", []) if isinstance(rows[0], dict) else rows[0]
                    table = doc.add_table(rows=len(rows), cols=len(first_row))
                    for i, row_data in enumerate(rows):
                        cells = row_data.get("cells", []) if isinstance(row_data, dict) else row_data
                        for j, cell_text in enumerate(cells):
                            if j < len(table.columns):
                                table.cell(i, j).text = str(cell_text)

        doc.save(file_path)

    def preview_patch(self, json_content: dict, patch: dict) -> dict:
        if patch.get("operation_type") not in ("replace_text", "modify_docx_paragraph"):
            raise ValueError("DOCX 补丁仅支持 replace_text 操作类型")

        blocks = _resolve_blocks(json_content)
        target_id = None
        for item in blocks:
            if item.get("type") in ("paragraph", "heading") and item.get("text", "").strip():
                target_id = item.get("id")
                break

        return {
            "preview_passed": True,
            "target_id": target_id,
            "risk_level": "medium",
            "style_loss_risk": True,
        }
