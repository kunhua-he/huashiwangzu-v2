from app.middleware.auth import require_permission
from app.models.user import User
from app.schemas.common import ApiResponse
from app.services.module_registry import register_capability
from app.services.parser_resource_diagnostics import (
    build_resource_diagnostic,
    store_extracted_resources_with_diagnostics,
)
from app.services.uploaded_file_runner import run_uploaded_file_capability
from fastapi import APIRouter, Depends
from pydantic import BaseModel

router = APIRouter(prefix="/api/docx-parser", tags=["docx-parser"])


class ParseRequest(BaseModel):
    file_id: int


async def _parse(params: dict, caller: str) -> dict:
    import base64

    from docx import Document as DocxDocument

    allowed = {"docx"}

    def parse_file(file_id, _file, full_path, _ext):
        doc = DocxDocument(str(full_path))
        blocks = []
        resources = []
        resource_diagnostics = []
        resource_counter = 0

        for para in doc.paragraphs:
            text = "\n".join(line.rstrip() for line in para.text.splitlines()).strip()
            if not text:
                continue
            style_name = str(para.style.name) if para.style else ""
            block_type = "heading" if ("heading" in style_name.lower() or "标题" in style_name) else "paragraph"
            blocks.append({"type": block_type, "text": text, "page": None, "resource_ref": None})

        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            table_text = "\n".join(rows)
            if table_text.strip():
                blocks.append({"type": "table", "text": table_text, "page": None, "resource_ref": None})

        for rel in doc.part.rels.values():
            if "image" in str(rel.reltype or "").lower():
                resource_counter += 1
                img_bytes = b""
                target_ref = str(getattr(rel, "target_ref", "") or "")
                mime_type = "image/png"
                extract_diagnostic_recorded = False
                try:
                    target_part = rel.target_part
                    img_bytes = target_part.blob
                    mime_type = getattr(target_part, "content_type", None) or "image/png"
                except Exception as exc:
                    extract_diagnostic_recorded = True
                    resource_diagnostics.append(build_resource_diagnostic(
                        parser="docx-parser",
                        stage="extract",
                        status="degraded",
                        code="resource_extract_failed",
                        message="Failed to extract DOCX embedded image bytes.",
                        resource={
                            "id": resource_counter,
                            "type": "image",
                            "filename": target_ref.split("/")[-1] if "/" in target_ref else "image.png",
                            "mime_type": mime_type,
                            "description": f"DOCX embedded image ({target_ref})",
                        },
                        error=exc,
                    ))
                blocks.append({"type": "image", "text": "", "page": None, "resource_ref": resource_counter})
                resources.append({
                    "id": resource_counter,
                    "type": "image",
                    "mime_type": mime_type,
                    "filename": target_ref.split("/")[-1] if "/" in target_ref else "image.png",
                    "description": f"DOCX embedded image ({target_ref})",
                    "_resource_diagnostic_recorded": extract_diagnostic_recorded,
                    "_bytes_b64": base64.b64encode(img_bytes).decode("ascii") if img_bytes else "",
                })

        return {
            "file_id": file_id,
            "format": "docx",
            "blocks": blocks,
            "resources": resources,
            "resource_diagnostics": resource_diagnostics,
        }

    result = await run_uploaded_file_capability(params, caller, allowed, parse_file)
    return await store_extracted_resources_with_diagnostics(result, caller=caller, parser="docx-parser")


@router.get("/health")
async def health():
    return ApiResponse(data={"module": "docx-parser", "status": "ok"})


@router.post("/parse")
async def call_parse(payload: ParseRequest, user: User = Depends(require_permission("viewer"))):
    result = await _parse({"file_id": payload.file_id}, f"user:{user.id}")
    return ApiResponse(data=result)


register_capability(
    "docx-parser", "parse", _parse,
    description="Parse DOCX files into unified content blocks",
    brief="解析 DOCX 文档",
    parameters={"file_id": {"type": "int"}},
    min_role="viewer",
)
