
"""Sandbox test for docx-parser module."""
import sys
from pathlib import Path
SAMPLE = Path(__file__).resolve().parent / "samples" / "sample.docx"
if not SAMPLE.exists():
    print("ERROR: sample.docx not found"); sys.exit(1)
from docx import Document as DocxDocument


def parse_docx(path):
    doc = DocxDocument(str(path))
    blocks, resources, rc = [], [], 0
    for para in doc.paragraphs:
        t = "\n".join(l.rstrip() for l in para.text.splitlines()).strip()
        if not t: continue
        style = str(para.style.name) if para.style else ""
        bt = "标题" if ("heading" in style.lower() or "标题" in style) else "段落"
        blocks.append({"type": bt, "text": t, "page": None, "resource_ref": None})
    for table in doc.tables:
        rows = [" | ".join(c.text.strip() for c in r.cells) for r in table.rows]
        tt = "\n".join(rows)
        if tt.strip(): blocks.append({"type": "表格", "text": tt, "page": None, "resource_ref": None})
    for rel in doc.part.rels.values():
        if "image" in str(rel.reltype or "").lower():
            rc += 1
            blocks.append({"type": "图片", "text": "", "page": None, "resource_ref": rc})
            resources.append({"id": rc, "type": "图片", "file_storage_id": None, "text_desc": f"DOCX image ({rel.target_ref})"})
    return {"file_id": 0, "format": "docx", "blocks": blocks, "resources": resources}


def validate(result):
    assert result["format"] == "docx"
    for b in result["blocks"]: assert all(k in b for k in ("type","text","page","resource_ref"))
    print("  Validation PASS (%d blocks, %d resources)" % (len(result["blocks"]), len(result["resources"])))


def main():
    print("="*60); print("docx-parser sandbox test"); print("="*60)
    result = parse_docx(SAMPLE)
    for b in result["blocks"]:
        print("    [%s] %s" % (b["type"], b["text"][:70]))
    validate(result)
    print("PASS: docx-parser sandbox test")

if __name__ == "__main__": main()
