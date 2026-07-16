from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path("/Users/hekunhua/Documents/Hermes工作区/竞品分析/直播文案_01版.md")
OUTPUT = Path("/Users/hekunhua/Documents/Hermes工作区/竞品分析/全肌品牌招商直播_精校原始文案_主播版.docx")

FONT_CN = "Microsoft YaHei"
FONT_EN = "Aptos"
INK = "1D1D1F"
MUTED = "6E6B66"
ORANGE = "F15A3A"
GREEN = "28776D"
BLUE = "3778A8"
PALE_ORANGE = "FCE4D6"
PALE_YELLOW = "FFF2CC"
PALE_GREEN = "E2F0D9"
LINE = "D8D2C8"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: int = 6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color: str = ORANGE, size: int = 18):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)


def set_run_font(run, size: float, color: str = INK, bold: bool = False, italic: bool = False):
    run.font.name = FONT_EN
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT_CN)


def add_inline_runs(paragraph, text: str, size: float = 14, color: str = INK):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size, color, True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, size - 0.5, ORANGE, True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size, color)


def set_body_paragraph(paragraph, first_line: bool = True):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    fmt.space_after = Pt(7)
    fmt.space_before = Pt(0)
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    return run


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_EN
    normal.font.size = Pt(14)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    heading_specs = {
        "Title": (30, INK, True),
        "Subtitle": (15, MUTED, False),
        "Heading 1": (22, ORANGE, True),
        "Heading 2": (18, INK, True),
        "Heading 3": (15.5, GREEN, True),
        "Heading 4": (14.5, BLUE, True),
    }
    for name, (size, color, bold) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = FONT_EN
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    doc.styles["Heading 1"].paragraph_format.space_before = Pt(0)
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(18)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(14)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(8)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(10)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(6)


def configure_section(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.7)


def add_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("全肌品牌招商直播｜精校原始干净稿 · 主播版")
    set_run_font(run, 8.5, MUTED)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("第 ")
    set_run_font(run, 8.5, MUTED)
    run = add_field(p, "PAGE")
    set_run_font(run, 8.5, MUTED)
    run = p.add_run(" 页 / 共 ")
    set_run_font(run, 8.5, MUTED)
    run = add_field(p, "NUMPAGES")
    set_run_font(run, 8.5, MUTED)
    run = p.add_run(" 页")
    set_run_font(run, 8.5, MUTED)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(64)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("竞品直播原始干净稿")
    set_run_font(run, 13, ORANGE, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("全肌品牌招商直播")
    set_run_font(run, 30, INK, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("精校原始文案 · 主播阅读版")
    set_run_font(run, 18, GREEN, True)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Cm(13.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, PALE_ORANGE)
    set_cell_border(cell, PALE_ORANGE, 1)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("适合主播熟悉整场结构、逐段练习与直播提词")
    set_run_font(run, 13.5, INK, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run("来源：20.02.55 / 20.09.27 / 21.07.47 三段直播录屏")
    set_run_font(run, 10.5, MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("内容只做精校还原，不代表对效果、收益及皮肤原理的核验")
    set_run_font(run, 10.5, MUTED)

    doc.add_page_break()


def add_guide(doc: Document, source: str):
    p = doc.add_paragraph(style="Heading 1")
    p.add_run("主播阅读说明")

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    rows = [
        ("重点话术", "正文中以加粗显示，建议熟记表达逻辑，不必机械背字。"),
        ("现场互动", "橙色提示框用于扣1、反问、价值确认和成交推进。"),
        ("视频演示", "黄色提示框为播放项目视频时配合使用的讲解话术。"),
    ]
    for i, (left, right) in enumerate(rows):
        c1, c2 = table.rows[i].cells
        c1.width = Cm(3.7)
        c2.width = Cm(12.5)
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(c1, [PALE_GREEN, PALE_ORANGE, PALE_YELLOW][i])
        set_cell_shading(c2, "FFFFFF")
        set_cell_border(c1)
        set_cell_border(c2)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p1.add_run(left)
        set_run_font(r, 12.5, INK, True)
        p2 = c2.paragraphs[0]
        r = p2.add_run(right)
        set_run_font(r, 11.5, MUTED)

    doc.add_paragraph()
    p = doc.add_paragraph(style="Heading 2")
    p.add_run("内容目录")
    headings = re.findall(r"^## (.+)$", source, flags=re.MULTILINE)
    for idx, heading in enumerate(headings, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.space_after = Pt(7)
        r = p.add_run(f"{idx:02d}  ")
        set_run_font(r, 11.5, ORANGE, True)
        r = p.add_run(heading)
        set_run_font(r, 13.5, INK, True)

    doc.add_page_break()


def add_heading(doc: Document, level: int, text: str, first_major: bool):
    if level == 2:
        if not first_major:
            doc.add_page_break()
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Cm(16.5)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, PALE_ORANGE)
        set_cell_border(cell, PALE_ORANGE, 1)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(text)
        set_run_font(r, 21, ORANGE, True)
        doc.add_paragraph()
        return

    style = {3: "Heading 2", 4: "Heading 3"}.get(level, "Heading 4")
    p = doc.add_paragraph(style=style)
    add_inline_runs(p, text, 18 if level == 3 else 15.5, INK if level == 3 else GREEN)
    p.paragraph_format.keep_with_next = True


def add_quote(doc: Document, text: str):
    is_video = "视频演示" in text
    is_interaction = "现场互动" in text or "主播互动" in text
    fill = PALE_YELLOW if is_video else PALE_ORANGE if is_interaction else "F2F2F2"
    color = ORANGE if is_interaction else GREEN if is_video else MUTED
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.right_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.35
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color)
    add_inline_runs(p, text, 12.5, INK)


def add_dialogue(doc: Document, text: str):
    match = re.match(r"\*\*(门店老板|郭总)：\*\*\s*(.*)", text)
    if not match:
        return False
    speaker, body = match.groups()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.25)
    p.paragraph_format.right_indent = Cm(0.1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.4
    fill = "F2F2F2" if speaker == "门店老板" else PALE_GREEN
    accent = MUTED if speaker == "门店老板" else GREEN
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, accent)
    r = p.add_run(f"{speaker}：")
    set_run_font(r, 13, accent, True)
    add_inline_runs(p, body, 13, INK)
    return True


def convert_markdown(doc: Document, source: str):
    lines = source.splitlines()
    in_preamble = True
    first_major = True
    for line in lines:
        stripped = line.strip()
        if in_preamble:
            if stripped.startswith("## "):
                in_preamble = False
            else:
                continue

        if not stripped:
            continue
        if stripped == "---":
            continue

        heading = re.match(r"^(#{2,5})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            add_heading(doc, level, heading.group(2), first_major)
            if level == 2:
                first_major = False
            continue

        if stripped.startswith("> "):
            add_quote(doc, stripped[2:])
            continue

        if add_dialogue(doc, stripped):
            continue

        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.first_line_indent = Cm(-0.35)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.35
            add_inline_runs(p, bullet.group(1), 13.5, INK)
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.95)
            p.paragraph_format.first_line_indent = Cm(-0.72)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.35
            r = p.add_run(f"{numbered.group(1)}. ")
            set_run_font(r, 13.5, ORANGE, True)
            add_inline_runs(p, numbered.group(2), 13.5, INK)
            continue

        p = doc.add_paragraph()
        set_body_paragraph(p, True)
        add_inline_runs(p, stripped, 14, INK)


def build_document():
    source = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_header_footer(doc.sections[0])
    doc.core_properties.title = "全肌品牌招商直播｜精校原始文案·主播版"
    doc.core_properties.subject = "竞品招商直播主播学习稿"
    doc.core_properties.author = "基于用户提供的ASR精校稿生成"

    add_cover(doc)
    add_guide(doc, source)
    convert_markdown(doc, source)

    # Keep table rows and headings together where possible.
    for table in doc.tables:
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = OxmlElement("w:cantSplit")
            tr_pr.append(cant_split)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_document()
    print(f"created {path}")
